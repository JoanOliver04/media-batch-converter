from __future__ import annotations

import os
import tempfile
import unittest
import zipfile
from pathlib import Path
from tkinter import TclError, Tk
from unittest.mock import patch

from documents.conversion import choose_engine, convert_document
from documents.errors import DocumentError
from documents.libreoffice import LibreOfficeInfo
from documents.formats import builtin_supports, conversion_supported, format_from_path
from documents.security import inspect_source, sniff_kind
from documents.settings import DocumentSettings, validate_document_settings
from documents.textio import parse_markdown, render_markdown
from error_handling import describe_error
from i18n import DEFAULT_LANGUAGE, set_language
from presets import CUSTOM_PRESET_ID, DOCUMENT_PRESETS, SettingsStore, preset_by_id
from ui.document_panel import DocumentPanel


class FormatMatrixTests(unittest.TestCase):
    def test_suffixes_map_to_canonical_formats(self) -> None:
        self.assertEqual(format_from_path("informe.DOCX"), "DOCX")
        self.assertEqual(format_from_path("notes.markdown"), "MD")
        self.assertEqual(format_from_path("page.HTM"), "HTML")
        self.assertIsNone(format_from_path("clip.mp4"))

    def test_builtin_covers_text_office_and_pdf_pairs(self) -> None:
        self.assertTrue(builtin_supports("DOCX", "PDF"))
        self.assertTrue(builtin_supports("PDF", "DOCX"))
        self.assertTrue(builtin_supports("MD", "HTML"))
        self.assertTrue(builtin_supports("XLSX", "CSV"))
        self.assertFalse(builtin_supports("DOC", "PDF"))
        self.assertTrue(conversion_supported("DOC", "PDF", "automatic", True))
        self.assertFalse(conversion_supported("DOC", "PDF", "builtin", True))

    def test_automatic_engine_prefers_libreoffice_for_layout(self) -> None:
        self.assertEqual(
            choose_engine("DOCX", "PDF", DocumentSettings(), True), "libreoffice"
        )
        self.assertEqual(
            choose_engine("DOCX", "PDF", DocumentSettings(), False), "builtin"
        )
        self.assertEqual(
            choose_engine("MD", "HTML", DocumentSettings(), True), "builtin"
        )
        self.assertEqual(
            choose_engine("DOC", "PDF", DocumentSettings(engine="automatic"), True),
            "libreoffice",
        )
        with self.assertRaises(DocumentError):
            choose_engine("DOC", "PDF", DocumentSettings(engine="builtin"), True)


class SecurityTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_rejects_empty_files_and_symlinks_and_type_mismatch(self) -> None:
        empty = self.root / "empty.txt"
        empty.write_bytes(b"")
        with self.assertRaises(DocumentError):
            inspect_source(empty)

        spoofed = self.root / "not-a-word.docx"
        spoofed.write_text("plain text pretending", encoding="utf-8")
        with self.assertRaises(DocumentError):
            inspect_source(spoofed)

        target = self.root / "real.txt"
        target.write_text("hola", encoding="utf-8")
        link = self.root / "alias.txt"
        try:
            os.symlink(target, link)
        except OSError:
            self.skipTest("symlinks are not available")
        with self.assertRaises(DocumentError):
            inspect_source(link)

    def test_rejects_zip_bombs_before_opening_office_packages(self) -> None:
        bomb = self.root / "bomb.docx"
        with zipfile.ZipFile(bomb, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("huge.txt", "A" * 2_000_000)
        with zipfile.ZipFile(bomb) as archive:
            info = archive.getinfo("huge.txt")
            self.assertGreater(info.file_size / max(info.compress_size, 1), 100)
        with self.assertRaises(DocumentError):
            inspect_source(bomb)

    def test_stored_zip_members_are_not_treated_as_bombs(self) -> None:
        package = self.root / "stored.docx"
        with zipfile.ZipFile(package, "w", compression=zipfile.ZIP_STORED) as archive:
            archive.writestr("[Content_Types].xml", "<Types/>")
            archive.writestr("word/document.xml", "<w:document/>")
        self.assertEqual(inspect_source(package).value, "zip")

    def test_sniffs_pdf_and_text_headers(self) -> None:
        pdf = self.root / "a.pdf"
        pdf.write_bytes(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        text = self.root / "a.txt"
        text.write_text("contenido", encoding="utf-8")
        self.assertEqual(sniff_kind(pdf).value, "pdf")
        self.assertEqual(sniff_kind(text).value, "text")


class ConversionRoundTripTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.settings = DocumentSettings(engine="builtin", page_markers=False)

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_markdown_to_docx_to_text_keeps_structure(self) -> None:
        source = self.root / "notes.md"
        source.write_text(
            "# Titulo\n\nUn parrafo con tildes: canción.\n\n- uno\n- dos\n\n"
            "| A | B |\n| --- | --- |\n| 1 | 2 |\n",
            encoding="utf-8",
        )
        docx = self.root / "notes.docx"
        convert_document(source, docx, "DOCX", self.settings)
        self.assertTrue(docx.is_file())
        text = self.root / "notes.txt"
        convert_document(docx, text, "TXT", self.settings)
        contents = text.read_text(encoding="utf-8")
        self.assertIn("Titulo", contents)
        self.assertIn("cancion" in contents or "canción" in contents, {True})
        self.assertIn("uno", contents)

    def test_txt_to_pdf_to_txt_extracts_body(self) -> None:
        source = self.root / "memo.txt"
        source.write_text("Primera linea.\n\nSegunda linea.", encoding="utf-8")
        pdf = self.root / "memo.pdf"
        convert_document(source, pdf, "PDF", self.settings)
        self.assertTrue(pdf.read_bytes().startswith(b"%PDF"))
        extracted = self.root / "memo-back.txt"
        convert_document(pdf, extracted, "TXT", DocumentSettings(engine="builtin"))
        body = extracted.read_text(encoding="utf-8")
        self.assertIn("Primera linea", body)
        self.assertIn("Segunda linea", body)

    def test_word_docx_keeps_images_and_header_in_pdf(self) -> None:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image

        from documents.conversion import read_document
        from documents.model import BlockKind

        picture = self.root / "marca.png"
        Image.new("RGB", (48, 24), "red").save(picture)
        source = self.root / "informe.docx"
        document = Document()
        document.sections[0].header.paragraphs[0].text = "Acme S.L."
        document.add_paragraph("Informe con figura")
        document.add_picture(str(picture), width=Inches(1.2))
        document.save(str(source))

        model = read_document(source, "DOCX", self.settings)
        self.assertEqual(model.header, "Acme S.L.")
        self.assertTrue(any(block.kind is BlockKind.IMAGE for block in model.blocks))

        pdf = self.root / "informe.pdf"
        convert_document(source, pdf, "PDF", self.settings)
        self.assertTrue(pdf.read_bytes().startswith(b"%PDF"))
        self.assertGreater(pdf.stat().st_size, 1_500)

        back = self.root / "informe-back.docx"
        convert_document(source, back, "DOCX", self.settings)
        restored = read_document(back, "DOCX", self.settings)
        self.assertTrue(any(block.kind is BlockKind.IMAGE for block in restored.blocks))

    def test_automatic_falls_back_to_builtin_when_libreoffice_fails(self) -> None:
        source = self.root / "letter.docx"
        from docx import Document

        document = Document()
        document.add_paragraph("Carta de presentacion.")
        document.save(str(source))
        output = self.root / "letter.pdf"
        office = LibreOfficeInfo(Path("C:/missing/soffice.exe"), "LibreOffice 0")
        with patch(
            "documents.conversion.convert_with_libreoffice",
            side_effect=DocumentError("LibreOffice crashed"),
        ):
            outcome = convert_document(
                source,
                output,
                "PDF",
                DocumentSettings(engine="automatic"),
                office=office,
            )
        self.assertEqual(outcome.engine, "builtin")
        self.assertTrue(output.is_file())
        self.assertTrue(output.read_bytes().startswith(b"%PDF"))
        self.assertTrue(any("LibreOffice" in warning for warning in outcome.warnings))

    def test_forced_libreoffice_does_not_fallback(self) -> None:
        source = self.root / "letter.docx"
        from docx import Document

        document = Document()
        document.add_paragraph("x")
        document.save(str(source))
        office = LibreOfficeInfo(Path("C:/missing/soffice.exe"), "LibreOffice 0")
        with patch(
            "documents.conversion.convert_with_libreoffice",
            side_effect=DocumentError("LibreOffice crashed"),
        ):
            with self.assertRaises(DocumentError):
                convert_document(
                    source,
                    self.root / "letter.pdf",
                    "PDF",
                    DocumentSettings(engine="libreoffice"),
                    office=office,
                )

    def test_csv_xlsx_round_trip(self) -> None:
        source = self.root / "data.csv"
        source.write_text("nombre,valor\nalfa,1\nbeta,2\n", encoding="utf-8")
        xlsx = self.root / "data.xlsx"
        convert_document(source, xlsx, "XLSX", self.settings)
        back = self.root / "data-back.csv"
        convert_document(xlsx, back, "CSV", self.settings)
        rows = back.read_text(encoding="utf-8-sig")
        self.assertIn("alfa", rows)
        self.assertIn("beta", rows)

    def test_html_to_markdown(self) -> None:
        source = self.root / "page.html"
        source.write_text(
            "<html><head><title>Guia</title></head><body>"
            "<h1>Guia</h1><p>Paso uno.</p><ul><li>A</li></ul></body></html>",
            encoding="utf-8",
        )
        markdown = self.root / "page.md"
        convert_document(source, markdown, "MD", self.settings)
        rendered = markdown.read_text(encoding="utf-8")
        self.assertIn("# Guia", rendered)
        self.assertIn("Paso uno.", rendered)
        self.assertIn("- A", rendered)

    def test_same_format_pdf_is_rewritten_not_reflowed(self) -> None:
        source = self.root / "origin.txt"
        source.write_text("Conservar.", encoding="utf-8")
        pdf = self.root / "origin.pdf"
        convert_document(source, pdf, "PDF", self.settings)
        copy = self.root / "copy.pdf"
        convert_document(pdf, copy, "PDF", self.settings)
        self.assertTrue(copy.read_bytes().startswith(b"%PDF"))

    def test_encrypted_pdf_fails_clearly(self) -> None:
        from pypdf import PdfWriter

        pdf = self.root / "secret.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.encrypt("password")
        with pdf.open("wb") as handle:
            writer.write(handle)
        with self.assertRaises(DocumentError) as raised:
            convert_document(pdf, self.root / "out.txt", "TXT", self.settings)
        self.assertIn("cifrado", str(raised.exception).casefold())

    def test_markdown_parser_skips_separator_rows(self) -> None:
        blocks = parse_markdown("| A | B |\n| --- | --- |\n| 1 | 2 |\n")
        tables = [block for block in blocks if block.rows]
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0].rows[0], ("A", "B"))
        self.assertEqual(len(tables[0].rows), 2)
        model_text = render_markdown(
            type("M", (), {"title": None, "blocks": tuple(blocks)})()
        )
        self.assertIn("| A | B |", model_text)


class SettingsAndPanelTests(unittest.TestCase):
    def setUp(self) -> None:
        self.addCleanup(set_language, DEFAULT_LANGUAGE)
        self.temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary.cleanup)
        previous = os.environ.get("APPDATA")
        os.environ["APPDATA"] = self.temporary.name
        self.addCleanup(self._restore, previous)
        try:
            self.root = Tk()
            self.root.withdraw()
        except TclError as error:
            self.skipTest(f"Tk unavailable: {error}")
        self.addCleanup(self.root.destroy)
        with patch("ui.document_panel.resolve_libreoffice", return_value=None):
            self.panel = DocumentPanel(self.root, self.root)

    @staticmethod
    def _restore(previous: str | None) -> None:
        if previous is None:
            os.environ.pop("APPDATA", None)
        else:
            os.environ["APPDATA"] = previous

    def test_invalid_settings_are_rejected(self) -> None:
        with self.assertRaises(ValueError):
            validate_document_settings(DocumentSettings(page_size="legal"))
        with self.assertRaises(ValueError):
            validate_document_settings(DocumentSettings(engine="word"))

    def test_preset_applies_and_manual_change_is_custom(self) -> None:
        self.panel.apply_document_preset_id("document_plain_text")
        self.assertEqual(self.panel.output_format.get(), "TXT")
        self.assertEqual(self.panel.current_document_settings().engine, "builtin")
        self.panel.output_format.set("PDF")
        self.assertEqual(self.panel.preset_display.get(), "Personalizado")
        self.assertEqual(
            self.panel.conversion_options()["document_preset"], CUSTOM_PRESET_ID
        )

    def test_filename_preview_uses_document_extension(self) -> None:
        source = Path(self.temporary.name) / "Informe Final.docx"
        source.write_bytes(b"PK\x03\x04preview")
        self.panel.selection.set(str(source))
        self.panel.output_format.set("PDF")
        self.panel.normalize_filenames.set(True)
        self.panel.update_output_name_preview()
        self.assertEqual(
            self.panel.output_name_preview.get(),
            "Nombre de salida: informe_final.pdf",
        )

    def test_single_file_validation_rejects_impossible_pairs(self) -> None:
        source = Path(self.temporary.name) / "legacy.doc"
        source.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 32)
        self.panel.selection.set(str(source))
        self.panel.output_format.set("PDF")
        self.panel.document_engine.set("Integrado")
        self.assertIsNotNone(self.panel.validate_start())

    def test_document_error_is_shown_as_its_own_message(self) -> None:
        description = describe_error(DocumentError("El PDF está cifrado."))
        self.assertEqual(description.message, "El PDF está cifrado.")

    def test_document_presets_are_registered(self) -> None:
        self.assertGreaterEqual(len(DOCUMENT_PRESETS), 5)
        preset = preset_by_id("document_pdf_archive")
        self.assertIsNotNone(preset)
        self.assertEqual(preset.output_format, "PDF")
        store = SettingsStore(Path(self.temporary.name) / "settings.json")
        store.save_last_document_preset("document_html_export")
        self.assertEqual(store.load_last_document_preset(), "document_html_export")
        store.save_last_document_preset("missing")
        self.assertEqual(store.load_last_document_preset(), CUSTOM_PRESET_ID)


if __name__ == "__main__":
    unittest.main()
